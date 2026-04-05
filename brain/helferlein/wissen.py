"""Wissen-Helferlein -- beobachtet Ordner und ingestiert neue Dokumente.

Scannt konfigurierte "Watch Folders" nach neuen oder geaenderten
Dokumenten (.docx, .md, .txt, .pdf) und konvertiert sie in die
Knowledge Base als Markdown. Dadurch lernt geofrey automatisch
wenn Slavko neue Kurs-Materialien, Recherchen oder Notizen ablegt.

Die konvertierten Dateien landen in knowledge-base/business/ (oder
einem anderen konfigurierten Unterordner) und werden beim naechsten
Lauf vom Knowledge Hub in ChromaDB indexiert.
"""

import hashlib
import json
import logging
from datetime import datetime
from pathlib import Path

from brain.helferlein import register
from brain.proposals import create_proposal, has_pending_proposal

logger = logging.getLogger("geofrey.helferlein.wissen")

# Watch folders: path -> knowledge-base subdirectory
WATCH_FOLDERS = {
    "~/Code/FBB-SlavkoKlincov": {
        "kb_subdir": "business",
        "beschreibung": "FBB Kurs: Leadgenerierung, No-Brainer Offers, Zielgruppen, Pricing",
    },
}

# Supported file types
SUPPORTED_EXTENSIONS = {".docx", ".md", ".txt"}

# Hash cache to detect changes
HASH_CACHE_PATH = Path.home() / ".knowledge" / "wissen_hashes.json"


def _load_hash_cache() -> dict[str, str]:
    """Load file hash cache."""
    if HASH_CACHE_PATH.exists():
        try:
            return json.loads(HASH_CACHE_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def _save_hash_cache(cache: dict[str, str]) -> None:
    """Save file hash cache."""
    HASH_CACHE_PATH.parent.mkdir(parents=True, exist_ok=True)
    HASH_CACHE_PATH.write_text(json.dumps(cache, indent=2), encoding="utf-8")


def _file_hash(path: Path) -> str:
    """Compute hash of file content + mtime."""
    stat = path.stat()
    return hashlib.md5(f"{path}:{stat.st_size}:{stat.st_mtime}".encode()).hexdigest()


def _docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed. Run: pip install python-docx")
        return ""

    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning(f"Could not read {path.name}: {e}")
        return ""

    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        # Convert heading styles to markdown
        style = para.style.name if para.style else ""
        if "Heading 1" in style or "heading 1" in style:
            lines.append(f"# {text}")
        elif "Heading 2" in style or "heading 2" in style:
            lines.append(f"## {text}")
        elif "Heading 3" in style or "heading 3" in style:
            lines.append(f"### {text}")
        elif "List" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Also extract tables
    for table in doc.tables:
        lines.append("")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines)


def _convert_file(source: Path) -> str:
    """Convert a file to markdown text based on extension."""
    ext = source.suffix.lower()
    if ext == ".docx":
        return _docx_to_markdown(source)
    elif ext in (".md", ".txt"):
        return source.read_text(encoding="utf-8")
    return ""


def _ingest_folder(
    folder_path: Path,
    kb_subdir: str,
    beschreibung: str,
    hash_cache: dict[str, str],
) -> tuple[int, list[str]]:
    """Ingest all supported files from a folder into the knowledge base.

    Returns (count_new, list_of_new_filenames).
    """
    geofrey_root = Path(__file__).parent.parent.parent
    kb_dir = geofrey_root / "knowledge-base" / kb_subdir
    kb_dir.mkdir(parents=True, exist_ok=True)

    new_files = []
    count = 0

    for source_file in sorted(folder_path.iterdir()):
        if source_file.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if source_file.name.startswith(".") or source_file.name.startswith("~"):
            continue

        # Check if file changed since last scan
        current_hash = _file_hash(source_file)
        cache_key = str(source_file)
        if hash_cache.get(cache_key) == current_hash:
            continue  # Unchanged

        # Convert to markdown
        markdown = _convert_file(source_file)
        if not markdown or len(markdown.strip()) < 50:
            continue

        # Write to knowledge base
        target_name = source_file.stem + ".md"
        target_path = kb_dir / target_name

        # Add frontmatter
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        content = (
            f"---\n"
            f"quelle: \"{source_file.name}\"\n"
            f"ordner: \"{folder_path.name}\"\n"
            f"konvertiert: \"{now}\"\n"
            f"beschreibung: \"{beschreibung}\"\n"
            f"---\n\n"
            f"{markdown}"
        )

        target_path.write_text(content, encoding="utf-8")
        hash_cache[cache_key] = current_hash
        new_files.append(source_file.name)
        count += 1
        logger.info(f"Ingested: {source_file.name} -> {kb_subdir}/{target_name}")

    return count, new_files


@register
class WissenHelferlein:
    """Watches folders for new documents and ingests them into the knowledge base."""

    name = "wissen"

    def run(self, config: dict) -> int:
        """Scan watch folders for new/changed documents. Returns proposals created."""
        hash_cache = _load_hash_cache()
        total_new = 0
        all_new_files: list[str] = []

        for folder_str, folder_config in WATCH_FOLDERS.items():
            folder_path = Path(folder_str).expanduser()
            if not folder_path.exists():
                logger.debug(f"Watch folder not found: {folder_str}")
                continue

            count, new_files = _ingest_folder(
                folder_path=folder_path,
                kb_subdir=folder_config["kb_subdir"],
                beschreibung=folder_config["beschreibung"],
                hash_cache=hash_cache,
            )
            total_new += count
            all_new_files.extend(new_files)

        _save_hash_cache(hash_cache)

        if not all_new_files:
            return 0

        # Create a proposal summarizing what was ingested
        if not has_pending_proposal("wissen", "Neue Dokumente"):
            create_proposal(
                helferlein="wissen",
                title=f"Neue Dokumente: {len(all_new_files)} File(s) in Knowledge Base",
                description=(
                    f"Folgende Dokumente wurden in die Knowledge Base aufgenommen:\n\n"
                    + "\n".join(f"- {f}" for f in all_new_files)
                    + f"\n\nDie Inhalte stehen geofrey jetzt fuer Recherche und "
                    f"Beratung zur Verfuegung."
                ),
                priority="low",
                action_type="notify",
            )

        logger.info(f"Ingested {total_new} new document(s) into knowledge base.")
        return 1 if all_new_files else 0
